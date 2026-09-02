#!/usr/bin/env python3
"""md -> docx 简易转换（标题/列表/行内加粗）。
用法: uv run --with python-docx python3 md2docx.py <in.md> <out.docx>
"""
import re, sys
from docx import Document
from docx.shared import Pt

def main():
    src = sys.argv[1] if len(sys.argv) > 1 else "in.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else "out.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    with open(src, encoding="utf-8") as f:
        lines = f.readlines()

    for raw in lines:
        line = raw.rstrip("\n")
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            # 行内 **bold** 分段处理
            texts = re.split(r"\*\*(.+?)\*\*", line)
            p = doc.add_paragraph()
            for i, t in enumerate(texts):
                if not t:
                    continue
                run = p.add_run(t)
                if i % 2 == 1:
                    run.bold = True
    doc.save(dst)
    print("saved", dst)

if __name__ == "__main__":
    main()
