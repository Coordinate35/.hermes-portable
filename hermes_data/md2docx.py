#!/usr/bin/env python3
"""md -> docx 简易转换（支持标题/列表/段落）"""
import re, sys
from docx import Document
from docx.shared import Pt

def main():
    src = "/home/coordinate35/hermes_data/魏俊杰-简历-v2.md"
    dst = "/home/coordinate35/hermes_data/魏俊杰-简历-v2.docx"
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)

    with open(src, encoding="utf-8") as f:
        lines = f.readlines()

    for raw in lines:
        line = raw.rstrip("\n")
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            # 行内 **bold** 简化处理
            texts = re.split(r"\*\*(.+?)\*\*", line)
            p = doc.add_paragraph()
            for i, t in enumerate(texts):
                if not t:
                    continue
                run = p.add_run(t)
                if i % 2 == 1:
                    run.bold = True
    doc.save(dst)
    print("saved", dst)

if __name__ == "__main__":
    main()
